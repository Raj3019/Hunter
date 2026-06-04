import re
from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from core.storage import create_signed_resume_url, resume_storage_path

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def create_tailored_resume_draft(
    db,
    user_id: str,
    match_data: dict[str, Any],
    resume_data: dict[str, Any],
    tailored: dict[str, Any],
) -> dict[str, Any]:
    _ensure_tailored_resumes_table(db)
    safe_tailored, validation = validate_tailored_resume(
        original_text=resume_data.get("raw_text", ""),
        resume_parsed=resume_data.get("parsed_data") or {},
        tailored=tailored,
    )
    version = _version()
    docx_bytes = build_tailored_resume_docx(
        resume_parsed=resume_data.get("parsed_data") or {},
        original_text=resume_data.get("raw_text", ""),
        job_data=match_data.get("jobs") or {},
        tailored=safe_tailored,
    )
    storage_path = _storage_path(
        user_id=user_id,
        match_id=match_data["id"],
        version=version,
        resume_data=resume_data,
    )

    db.storage.from_("resumes").upload(
        storage_path,
        docx_bytes,
        {"content-type": DOCX_MIME, "upsert": "true"},
    )
    stored_file_url = db.storage.from_("resumes").get_public_url(storage_path)
    signed_file_url = create_signed_resume_url(db, storage_path)

    payload = {
        "user_id": user_id,
        "match_id": match_data["id"],
        "source_resume_id": resume_data.get("id"),
        "status": "draft" if validation["ok"] else "failed_validation",
        "file_url": stored_file_url,
        "file_type": "docx",
        "version": version,
        "tailoring_json": safe_tailored,
        "validation_json": validation,
    }
    result = db.table("tailored_resumes").insert(payload).execute()
    row = result.data[0] if result.data else payload

    return {
        "id": row.get("id"),
        "status": row.get("status", payload["status"]),
        "file_url": signed_file_url or create_signed_resume_url(db, row.get("file_url", "")),
        "file_type": row.get("file_type", "docx"),
        "version": row.get("version", version),
        "tailoring": row.get("tailoring_json") or safe_tailored,
        "validation": row.get("validation_json") or validation,
    }


def validate_tailored_resume(
    original_text: str,
    resume_parsed: dict[str, Any],
    tailored: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]]:
    safe_tailored = deepcopy(tailored)
    original_norm = _normalize(original_text)
    original_skills = _string_list(resume_parsed.get("skills"))
    original_skill_norms = {_normalize(skill) for skill in original_skills}

    warnings: list[str] = []
    blocked_claims: list[str] = []
    removed_skills: list[str] = []
    valid_skills: list[str] = []

    for skill in _string_list(safe_tailored.get("reordered_skills")):
        skill_norm = _normalize(skill)
        if skill_norm in original_skill_norms or skill_norm in original_norm:
            valid_skills.append(skill)
        else:
            removed_skills.append(skill)

    if removed_skills:
        safe_tailored["reordered_skills"] = valid_skills
        warnings.append(
            "Removed skills not grounded in the uploaded resume: "
            + ", ".join(removed_skills)
        )

    candidate_years = _number(resume_parsed.get("total_experience_years"))
    if candidate_years is not None:
        combined_text = " ".join(
            [
                _text(safe_tailored.get("tailored_summary")),
                " ".join(_string_list(safe_tailored.get("highlighted_experience"))),
            ]
        )
        for years in _mentioned_years(combined_text):
            if years > candidate_years:
                blocked_claims.append(
                    f"Tailored draft claims {years} years of experience, but parsed resume has {candidate_years:g}."
                )

    model_warning = _text(safe_tailored.get("warnings"))
    if model_warning:
        warnings.append(model_warning)

    safe_tailored.setdefault("tailored_summary", "")
    safe_tailored.setdefault("reordered_skills", valid_skills)
    safe_tailored.setdefault("highlighted_experience", [])
    safe_tailored.setdefault("changes_made", [])
    safe_tailored["warnings"] = "; ".join(warnings)

    validation = {
        "ok": not blocked_claims,
        "blocked_claims": blocked_claims,
        "warnings": warnings,
        "removed_skills": removed_skills,
    }
    return safe_tailored, validation


def build_tailored_resume_docx(
    resume_parsed: dict[str, Any],
    original_text: str,
    job_data: dict[str, Any],
    tailored: dict[str, Any],
) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = None

    name = _text(resume_parsed.get("name"), "Candidate")
    document.add_heading(name, level=0)

    contact_parts = [
        _text(resume_parsed.get("email")),
        _text(resume_parsed.get("phone")),
        _text(resume_parsed.get("location")),
        _text(resume_parsed.get("linkedin_url")),
        _text(resume_parsed.get("github_url")),
    ]
    _add_paragraph(document, " | ".join(part for part in contact_parts if part))

    current_role = _text(resume_parsed.get("current_role"))
    years = resume_parsed.get("total_experience_years")
    if current_role or years not in ("", None):
        role_line = current_role
        if years not in ("", None):
            role_line = f"{role_line} | {years} years experience" if role_line else f"{years} years experience"
        _add_paragraph(document, role_line)

    _add_section(document, "Target Role")
    _add_paragraph(
        document,
        " - ".join(
            part
            for part in [
                _text(job_data.get("title")),
                _text(job_data.get("company")),
                _text(job_data.get("location")),
            ]
            if part
        ),
    )

    _add_section(document, "Professional Summary")
    _add_paragraph(
        document,
        _text(tailored.get("tailored_summary"))
        or _text(resume_parsed.get("summary"))
        or "Summary available in the uploaded resume.",
    )

    skills = _string_list(tailored.get("reordered_skills")) or _string_list(resume_parsed.get("skills"))
    if skills:
        _add_section(document, "Relevant Skills")
        _add_bullets(document, skills[:18])

    highlights = _string_list(tailored.get("highlighted_experience"))
    if highlights:
        _add_section(document, "Role-Aligned Experience Highlights")
        _add_bullets(document, highlights[:8])

    education = _text(resume_parsed.get("education"))
    education_details = _string_list(resume_parsed.get("education_details"))
    if education or education_details:
        _add_section(document, "Education")
        if education:
            _add_paragraph(document, education)
        _add_bullets(document, education_details[:4])

    certifications = _string_list(resume_parsed.get("certifications"))
    if certifications:
        _add_section(document, "Certifications")
        _add_bullets(document, certifications[:8])

    original_sections = _extract_original_resume_sections(original_text)
    if original_sections:
        _add_section(document, "Original Resume Content")
        for section in original_sections:
            heading = section["heading"]
            body = section["body"]
            if heading:
                document.add_heading(heading, level=2)
            for line in body[:14]:
                _add_paragraph(document, line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_section(document: Document, title: str) -> None:
    document.add_heading(title, level=1)


def _add_paragraph(document: Document, text: str) -> None:
    if text:
        document.add_paragraph(text)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        if item:
            document.add_paragraph(item, style="List Bullet")


def _version() -> str:
    return "tailored:" + datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def _storage_path(
    user_id: str,
    match_id: str,
    version: str,
    resume_data: dict[str, Any],
) -> str:
    safe_version = re.sub(r"[^A-Za-z0-9_.-]+", "-", version).strip("-")
    filename = _tailored_filename(resume_data)
    return f"{user_id}/tailored/{match_id}/{safe_version}/{filename}"


def _ensure_tailored_resumes_table(db) -> None:
    db.table("tailored_resumes").select("id").limit(1).execute()


def _mentioned_years(text: str) -> list[float]:
    matches = re.findall(r"\b(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\b", text, flags=re.IGNORECASE)
    return [float(value) for value in matches]


def _string_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_format_list_item(item) for item in value if _format_list_item(item)]
    if isinstance(value, str) and value.strip():
        return [part.strip() for part in value.split(",") if part.strip()]
    if isinstance(value, dict):
        formatted = _format_list_item(value)
        return [formatted] if formatted else []
    return []


def _format_list_item(value: Any) -> str:
    if isinstance(value, dict):
        preferred = [
            value.get("degree"),
            value.get("institution"),
            value.get("year"),
            value.get("field"),
            value.get("score"),
        ]
        parts = [_text(part) for part in preferred if _text(part)]
        if parts:
            return ", ".join(parts)
        return ", ".join(f"{key}: {_text(item)}" for key, item in value.items() if _text(item))
    return _text(value)


def _extract_original_resume_sections(original_text: str) -> list[dict[str, list[str] | str]]:
    lines = [line.strip(" -\t") for line in original_text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        return []

    headings = {
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "projects",
        "project",
        "education",
        "certifications",
        "achievements",
        "skills",
    }
    sections: list[dict[str, Any]] = []
    current = {"heading": "Resume Details", "body": []}

    for line in lines[:160]:
        normalized = re.sub(r"[^a-z ]+", "", line.lower()).strip()
        is_heading = normalized in headings or (
            len(line) <= 36
            and line.isupper()
            and len(line.split()) <= 4
        )
        if is_heading and current["body"]:
            sections.append(current)
            current = {"heading": line.title(), "body": []}
            continue
        if is_heading:
            current["heading"] = line.title()
            continue
        current["body"].append(line)

    if current["body"]:
        sections.append(current)

    return [section for section in sections if section["body"]][:8]


def _tailored_filename(resume_data: dict[str, Any]) -> str:
    storage_path = resume_storage_path(_text(resume_data.get("file_url"))) or _text(resume_data.get("file_url"))
    original_name = Path(str(storage_path).split("?", 1)[0]).name or "resume.pdf"
    if not original_name or "." not in original_name:
        original_name = "resume.pdf"

    stem = Path(original_name).stem
    safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "-", stem).strip("-") or "resume"
    return f"{safe_stem}-tailor.docx"


def _text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    if isinstance(value, str):
        return value.strip() or fallback
    return str(value).strip() or fallback


def _number(value: Any) -> float | None:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if number >= 0 else None


def _normalize(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", _text(value).lower())
